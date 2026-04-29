"""
cv_import_service.py — Pipeline d'ingestion et d'analyse unitaire d'un CV.

Ce module contient :
- _fetch_cv_content()
- _process_cv_core()

Il gère le téléchargement, l'analyse IA, la résolution d'identité, l'assignation
des compétences/missions et l'embedding pour un seul CV.
"""

import asyncio
import json
import logging
import os
import random
import re
import string
import time
import unicodedata
import urllib.parse
from typing import Any, List, Optional

import httpx
from fastapi import BackgroundTasks, HTTPException
from opentelemetry.propagate import inject
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy import delete as sa_delete

import database
from src.cvs.models import CVProfile
from src.cvs.schemas import CVImportStep, CVResponse
from src.cvs.task_state import task_state_manager
from metrics import CV_PROCESSING_TOTAL

from src.gemini_retry import generate_content_with_retry, embed_content_with_retry
from google import genai
from google.genai import types

from src.services.config import (
    COMPETENCIES_API_URL,
    ITEMS_API_URL,
    PROMPTS_API_URL,
    USERS_API_URL,
    _CV_CACHE,
)
from src.services.finops import log_finops
from src.services.utils import (
    _CV_RESPONSE_SCHEMA,
    _build_distilled_content,
    _coerce_to_str,
    build_taxonomy_context,
)
from datetime import datetime, timedelta, timezone

logger = logging.getLogger(__name__)


async def _fetch_cv_content(
    url: str,
    google_token: Optional[str] = None,
    file_type: str = "google_doc",
) -> str:
    """Download the CV content.

    - google_doc : export texte via Drive v3 /files/{id}/export?mimeType=text/plain
    - docx       : download binaire via /files/{id}?alt=media + extraction python-docx
    """
    parsed = urllib.parse.urlparse(url)
    hostname = parsed.hostname or ""

    if parsed.scheme not in ("http", "https"):
        raise HTTPException(status_code=400, detail="Invalid URL scheme")

    forbidden_hosts = ["localhost", "127.0.0.1", "0.0.0.0"]
    if hostname in forbidden_hosts or hostname.endswith(".local") or hostname.endswith("_api"):
        raise HTTPException(status_code=400, detail="Internal URLs are not allowed")

    # ── Branche DOCX ──────────────────────────────────────────────────────────
    if file_type == "docx":
        file_id_match = re.search(r"/file/d/([a-zA-Z0-9_-]+)", url)
        file_id = file_id_match.group(1) if file_id_match else None
        if not file_id:
            raise HTTPException(
                status_code=400,
                detail="DOCX URL invalide : impossible d'extraire le file_id depuis l'URL Drive.",
            )
        if not google_token:
            raise HTTPException(
                status_code=400,
                detail="DOCX Drive : google_access_token OAuth2 requis pour télécharger le fichier.",
            )
        download_url = f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"
        async with httpx.AsyncClient(timeout=60.0) as http_client:
            resp = await http_client.get(
                download_url,
                headers={"Authorization": f"Bearer {google_token}"},
                follow_redirects=True,
            )
            if resp.status_code in (401, 403):
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"Accès refusé pour le DOCX Drive (HTTP {resp.status_code}). "
                        "Vérifiez les scopes OAuth2 du Service Account (drive.readonly minimum requis)."
                    ),
                )
            resp.raise_for_status()

            from io import BytesIO
            import docx as python_docx
            doc = python_docx.Document(BytesIO(resp.content))

            # Extraction paragraphes (corps du CV)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Extraction tables (missions en tableau, expériences, compétences tabulées)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        paragraphs.append(row_text)

            text = "\n".join(paragraphs)
            if not text.strip():
                raise HTTPException(
                    status_code=422,
                    detail="Le fichier DOCX est vide ou illisible (aucun texte extrait).",
                )
            logger.info(f"[_fetch_cv_content] DOCX extrait — {len(text)} chars, file_id={file_id}")
            return text

    # ── Branche Google Doc natif ───────────────────────────────────────────────
    if "docs.google.com/document/d/" in url:
        doc_id = re.search(r"/d/([a-zA-Z0-9_-]+)", url)
        if doc_id:
            file_id = doc_id.group(1)
            if not google_token:
                logger.error(
                    f"[_fetch_cv_content] Aucun google_access_token pour le fichier Drive '{file_id}'. "
                    "L'API Drive v3 requiert un token OAuth2 valide (scope drive.readonly minimum)."
                )
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "Document Google Docs privé : un google_access_token OAuth2 est requis. "
                        "Vérifiez que drive_api transmet bien le token dans le payload Pub/Sub."
                    ),
                )
            export_url = f"https://www.googleapis.com/drive/v3/files/{file_id}/export"
            async with httpx.AsyncClient(timeout=30.0) as http_client:
                resp = await http_client.get(
                    export_url,
                    params={"mimeType": "text/plain"},
                    headers={"Authorization": f"Bearer {google_token}"},
                    follow_redirects=True,
                )
                if resp.status_code in (401, 403):
                    raise HTTPException(
                        status_code=400,
                        detail=(
                            f"Accès refusé par l'API Drive (HTTP {resp.status_code}). "
                            "Vérifiez les scopes OAuth2 du Service Account "
                            "(drive.readonly ou drive.file minimum requis)."
                        ),
                    )
                resp.raise_for_status()
                return resp.text

    req_headers = {}
    if google_token:
        req_headers["Authorization"] = f"Bearer {google_token}"

    async with httpx.AsyncClient(timeout=30.0) as http_client:
        resp = await http_client.get(url, headers=req_headers, follow_redirects=True)
        if resp.status_code in (401, 403):
            raise HTTPException(
                status_code=400,
                detail="Accès refusé. Vérifiez que le document est accessible.",
            )
        resp.raise_for_status()
        return resp.text


async def process_cv_core(
    url: str,
    google_access_token: Optional[str],
    source_tag: Optional[str],
    headers: dict,
    token_payload: dict,
    db: AsyncSession,
    auth_token: str = None,
    folder_name: Optional[str] = None,
    file_type: str = "google_doc",
    background_tasks: BackgroundTasks = None,
    genai_client=None,
) -> CVResponse:
    """
    Pipeline principal d'ingéstion d'un CV en 8 étapes séquentielles.
    Retourne un CVResponse enrichi avec les étapes (steps) et les warnings non-bloquants.
    """
    pipeline_steps: List[CVImportStep] = []
    pipeline_warnings: List[str] = []

    def _step_ok(step: str, label: str, duration_ms: int, detail: str = None) -> CVImportStep:
        s = CVImportStep(step=step, label=label, status="success", duration_ms=duration_ms, detail=detail)
        pipeline_steps.append(s)
        logger.info(f"[CV_STEP] {label} — OK", extra={"step": step, "duration_ms": duration_ms, "cv_url": url, "detail": detail})
        return s

    def _step_warn(step: str, label: str, duration_ms: int, detail: str = None) -> CVImportStep:
        s = CVImportStep(step=step, label=label, status="warning", duration_ms=duration_ms, detail=detail)
        pipeline_steps.append(s)
        pipeline_warnings.append(detail or label)
        logger.warning(f"[CV_STEP] {label} — WARN: {detail}", extra={"step": step, "duration_ms": duration_ms, "cv_url": url})
        return s

    def _step_error(step: str, label: str, duration_ms: int, detail: str = None) -> CVImportStep:
        s = CVImportStep(step=step, label=label, status="error", duration_ms=duration_ms, detail=detail)
        pipeline_steps.append(s)
        logger.error(f"[CV_STEP] {label} — ERROR: {detail}", extra={"step": step, "duration_ms": duration_ms, "cv_url": url})
        return s

    # ── Étape 1 : Téléchargement du document ─────────────────────────────────
    t0 = time.monotonic()
    try:
        logger.info(f"[CV_STEP] download — start (file_type={file_type})", extra={"step": "download", "cv_url": url})
        raw_text = await _fetch_cv_content(url, google_access_token, file_type=file_type)
        dur = int((time.monotonic() - t0) * 1000)
        raw_len = len(raw_text)
        if raw_len > 100000:
            warn_msg = f"Document tronqué : {raw_len} caractères → limité à 100000 pour l'analyse IA"
            _step_warn("download", "Téléchargement du document", dur, warn_msg)
        else:
            _step_ok("download", "Téléchargement du document", dur, f"{raw_len} caractères")
    except HTTPException as he:
        dur = int((time.monotonic() - t0) * 1000)
        _step_error("download", "Téléchargement du document", dur, he.detail)
        raise
    except Exception as e:
        dur = int((time.monotonic() - t0) * 1000)
        _step_error("download", "Téléchargement du document", dur, str(e))
        CV_PROCESSING_TOTAL.labels(status="failure").inc()
        raise HTTPException(status_code=400, detail=f"Failed downloading CV content: {e}")

    if not genai_client:
        raise HTTPException(status_code=500, detail="GenAI Client not configured.")

    # ── Étape 2 : Analyse IA — Extraction du profil ──────────────────────────
    t0 = time.monotonic()
    prompt = None
    if _CV_CACHE["prompt"]["expires"] > datetime.now(timezone.utc) and _CV_CACHE["prompt"]["value"]:
        prompt = _CV_CACHE["prompt"]["value"]
    else:
        try:
            async with httpx.AsyncClient() as http_client:
                res_prompt = await http_client.get(f"{PROMPTS_API_URL.rstrip('/')}/cv_api.extract_cv_info", headers=headers, timeout=5.0)
                res_prompt.raise_for_status()
                prompt = res_prompt.json()["value"]
                _CV_CACHE["prompt"]["value"] = prompt
                _CV_CACHE["prompt"]["expires"] = datetime.now(timezone.utc) + timedelta(minutes=5)
        except Exception as e:
            if os.path.exists("cv_api.extract_cv_info.txt"):
                with open("cv_api.extract_cv_info.txt", "r", encoding="utf-8") as f:
                    prompt = f.read()
            else:
                raise HTTPException(status_code=500, detail=f"Cannot fetch generic prompt: {e}")

    tree_context = ""
    if _CV_CACHE["tree_context"]["expires"] > datetime.now(timezone.utc) and _CV_CACHE["tree_context"]["value"]:
        tree_context = _CV_CACHE["tree_context"]["value"]
    else:
        try:
            async with httpx.AsyncClient() as http_client:
                items: list = []
                skip = 0
                page_size = 100
                while True:
                    page_res = await http_client.get(
                        f"{COMPETENCIES_API_URL.rstrip('/')}/",
                        params={"skip": skip, "limit": page_size},
                        headers=headers, timeout=5.0
                    )
                    if page_res.status_code != 200:
                        break
                    page_items = page_res.json().get('items', [])
                    items.extend(page_items)
                    if len(page_items) < page_size:
                        break
                    skip += page_size

                if items:
                    tree_context, nb_parents, nb_leaves = build_taxonomy_context(items)
                    _CV_CACHE["tree_context"]["value"] = tree_context
                    _CV_CACHE["tree_context"]["expires"] = datetime.now(timezone.utc) + timedelta(minutes=5)
                    _CV_CACHE["tree_items"]["value"] = items
                    _CV_CACHE["tree_items"]["expires"] = datetime.now(timezone.utc) + timedelta(minutes=5)
        except Exception as e:
            logger.warning(f"Failed to fetch competencies tree for context: {e}")

    final_prompt = prompt + tree_context

    try:
        response = await generate_content_with_retry(
            genai_client,
            model=os.getenv("GEMINI_CV_MODEL", os.getenv("GEMINI_MODEL")),
            contents=[final_prompt, f"RESUME:\n{raw_text[:100000]}"],
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=_CV_RESPONSE_SCHEMA,
            )
        )
        parsed_data = response.text
        structured_cv = json.loads(parsed_data)

        user_caller = token_payload.get("sub", "unknown")
        safe_meta = response.usage_metadata if hasattr(response, 'usage_metadata') else None
        await log_finops(user_caller, "analyze_cv", os.getenv("GEMINI_CV_MODEL", os.getenv("GEMINI_MODEL")), safe_meta, {"cv_url": url}, auth_token=auth_token)

        if not structured_cv.get("is_cv", False):
            dur = int((time.monotonic() - t0) * 1000)
            _step_error("llm_parse", "Analyse IA — Extraction du profil", dur, "Document non reconnu comme un CV")
            raise HTTPException(status_code=400, detail="Not a CV: The document does not appear to be a resume.")

        nb_competencies = len(structured_cv.get("competencies", []))
        nb_missions = len(structured_cv.get("missions", []))
        summary_val = structured_cv.get("summary", "")
        dur = int((time.monotonic() - t0) * 1000)

        llm_detail = f"{nb_competencies} compétences, {nb_missions} missions détectées"
        if nb_competencies == 0:
            warn_cv = "Aucune compétence extraite par l'IA — vérifiez la qualité du document"
            _step_warn("llm_parse", "Analyse IA — Extraction du profil", dur, llm_detail)
            pipeline_warnings.append(warn_cv)
        elif not summary_val or summary_val.strip() == "":
            warn_cv = "Résumé de profil vide — l'IA n'a pas pu générer de synthèse"
            _step_warn("llm_parse", "Analyse IA — Extraction du profil", dur, llm_detail)
            pipeline_warnings.append(warn_cv)
        else:
            _step_ok("llm_parse", "Analyse IA — Extraction du profil", dur, llm_detail)

    except HTTPException:
        raise
    except Exception as e:
        dur = int((time.monotonic() - t0) * 1000)
        _step_error("llm_parse", "Analyse IA — Extraction du profil", dur, str(e))
        CV_PROCESSING_TOTAL.labels(status="failure").inc()
        raise HTTPException(status_code=500, detail=f"LLM Parsing failed: {e}")

    # ── Étape 3 : Résolution d'identité ──────────────────────────────────────
    t0 = time.monotonic()

    def sanitize_field(val: Any) -> Optional[str]:
        if val is None: return None
        s = str(val).strip()
        clean_s = s.lower().strip(",").strip()
        if clean_s in ("null", "none", "", "unknown", "n/a", "na", "nil"): return None
        return s

    def normalize_str(s: str) -> str:
        if not s: return ""
        return unicodedata.normalize('NFKD', s).encode('ASCII', 'ignore').decode('utf-8').lower()

    NAME_REGEX = r"^[A-Za-zÀ-ÿ\s\-\']+$"
    EMAIL_REGEX = r'^[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+$'

    def is_valid_name(n: Optional[str]) -> bool:
        return bool(n and re.match(NAME_REGEX, n))
    def is_valid_email(e: Optional[str]) -> bool:
        return bool(e and re.match(EMAIL_REGEX, e))

    raw_email = sanitize_field(structured_cv.get("email"))
    llm_first_name = sanitize_field(structured_cv.get("first_name"))
    llm_last_name = sanitize_field(structured_cv.get("last_name"))
    is_anonymous = structured_cv.get("is_anonymous", False)
    trigram = sanitize_field(structured_cv.get("trigram"))

    if llm_first_name and not is_valid_name(llm_first_name):
        llm_first_name = None
    if llm_last_name and not is_valid_name(llm_last_name):
        llm_last_name = None

    folder_first_name: Optional[str] = None
    folder_last_name: Optional[str] = None

    if folder_name and folder_name.strip():
        parts = folder_name.strip().split(None, 1)
        if len(parts) == 2:
            folder_first_name = sanitize_field(parts[0])
            folder_last_name = sanitize_field(parts[1])
            if not is_valid_name(folder_first_name): folder_first_name = None
            if not is_valid_name(folder_last_name): folder_last_name = None

    first_name = folder_first_name or llm_first_name
    last_name = folder_last_name or llm_last_name

    if folder_first_name and folder_last_name and llm_first_name and llm_last_name:
        fn_match = normalize_str(folder_first_name) == normalize_str(llm_first_name)
        ln_match = normalize_str(folder_last_name) == normalize_str(llm_last_name)
        
        if not (fn_match and ln_match):
            folder_full = normalize_str(folder_first_name + folder_last_name).replace(' ', '').replace('-', '')
            llm_full = normalize_str(llm_first_name + llm_last_name).replace(' ', '').replace('-', '')
            
            if folder_full == llm_full:
                # Same identity, but the LLM did a better job splitting first_name and last_name 
                # (e.g. for compound names like "Jean Paul Belmondo" where folder blindly splits at first space)
                first_name, last_name = llm_first_name, llm_last_name
            else:
                warn_folder = f"⚠️ Divergence d'identité — Dossier: '{folder_first_name} {folder_last_name}' / LLM: '{llm_first_name} {llm_last_name}'."
                _step_warn("folder_identity", "Résolution identité — Divergence dossier vs LLM", 0, warn_folder)
                first_name, last_name = folder_first_name, folder_last_name

    if not is_valid_name(first_name): first_name = None
    if not is_valid_name(last_name): last_name = None

    if not is_valid_email(raw_email):
        if first_name and last_name:
            email = f"{normalize_str(first_name).replace(' ', '').replace('-', '')}.{normalize_str(last_name).replace(' ', '').replace('-', '')}@zenika.com"
            pipeline_warnings.append(f"Email absent ou invalide dans le CV — email généré : {email}")
        else:
            is_anonymous = True
            trigram = trigram or ''.join(random.choices(string.ascii_uppercase, k=3))
            first_name, last_name = "Anon", trigram
            email = f"anon.{trigram.lower()}@anonymous.zenika.com"
            pipeline_warnings.append("Identité introuvable dans le CV — profil anonymisé automatiquement")
    else:
        email = raw_email

    ext_full_norm = f"{normalize_str(first_name)} {normalize_str(last_name)}"

    async with httpx.AsyncClient(timeout=30.0) as http_client:
        user_id = None
        importer_id = None

        try:
            existing_cv = (await db.execute(select(CVProfile).where(CVProfile.source_url == url))).scalars().first()
        except AttributeError:
            existing_cv = (await db.execute(select(CVProfile).where(CVProfile.source_url == url))).first()
        if existing_cv and existing_cv.user_id:
            user_id = existing_cv.user_id

        if not user_id and folder_first_name and folder_last_name:
            folder_search_q = f"{folder_first_name} {folder_last_name}"
            fn_res = await http_client.get(f"{USERS_API_URL.rstrip('/')}/search", params={"query": folder_search_q, "limit": 10}, headers=headers)
            if fn_res.status_code == 200:
                for u in fn_res.json().get("items", []):
                    if normalize_str(u.get("first_name")) == normalize_str(folder_first_name) and normalize_str(u.get("last_name")) == normalize_str(folder_last_name):
                        user_id = u["id"]
                        break

        if not user_id and email:
            search_res = await http_client.get(f"{USERS_API_URL.rstrip('/')}/search", params={"query": email, "limit": 10}, headers=headers)
            if search_res.status_code == 200:
                for u in search_res.json().get("items", []):
                    if u.get("email", "").lower() == email.lower():
                        u_full_norm = normalize_str(u.get("full_name") or f"{u.get('first_name')} {u.get('last_name')}")
                        if ext_full_norm and ext_full_norm not in u_full_norm and u_full_norm not in ext_full_norm:
                            continue
                        user_id = u["id"]
                        break

        if not user_id and first_name and last_name:
            search_q = f"{first_name} {last_name}"
            name_res = await http_client.get(f"{USERS_API_URL.rstrip('/')}/search", params={"query": search_q, "limit": 10}, headers=headers)
            if name_res.status_code == 200:
                for u in name_res.json().get("items", []):
                    if normalize_str(u.get("first_name")) == normalize_str(first_name) and normalize_str(u.get("last_name")) == normalize_str(last_name):
                        user_id = u["id"]
                        break

        importer_username = token_payload.get("sub")
        if importer_username:
            importer_res = await http_client.get(f"{USERS_API_URL.rstrip('/')}/search", params={"query": importer_username, "limit": 10}, headers=headers)
            if importer_res.status_code == 200:
                for u in importer_res.json().get("items", []):
                    if u.get("username", "").lower() == importer_username.lower():
                        importer_id = u["id"]
                        break

        filename = os.path.basename(url).lower()
        if not is_anonymous and any(x in filename for x in ["annonym", "anon", "abc"]):
            is_anonymous = True
            if first_name != "Anon":
                trigram = trigram or ''.join(random.choices(string.ascii_uppercase, k=3))
                first_name, last_name = "Anon", trigram
                email = f"anon.{trigram.lower()}@anonymous.zenika.com"

        if user_id and is_anonymous:
            # Règle Zenika : le nom du dossier fait foi. Si un nom de dossier est fourni,
            # on lève l'anonymat pour permettre l'attachement au compte nominatif.
            if folder_first_name and folder_last_name:
                is_anonymous = False
            else:
                user_info_res = await http_client.get(f"{USERS_API_URL.rstrip('/')}/{user_id}", headers=headers)
                if user_info_res.status_code == 200:
                    user_data = user_info_res.json()
                    if not user_data.get("is_anonymous", False):
                        await task_state_manager.update_progress(new_log=f"🛡️ CV anonyme détecté sur un compte réel (User {user_id}). DÉTACHEMENT du profil.")
                        user_id = None

        if not user_id:
            safe_fn = first_name or "u"
            safe_ln = last_name or f"user{random.randint(1000,9999)}"
            new_u = {
                "username": f"{safe_fn[0].lower()}{safe_ln.lower().replace(' ', '')}{random.randint(100, 999)}",
                "email": email,
                "first_name": first_name,
                "last_name": last_name,
                "full_name": f"{first_name or ''} {last_name or ''}".strip() or "Unknown User",
                "password": "zenikacv123",
                "is_anonymous": is_anonymous
            }
            create_res = await http_client.post(f"{USERS_API_URL.rstrip('/')}/", json=new_u, headers=headers)
            if create_res.status_code == 409 or (create_res.status_code >= 400 and "already exists" in create_res.text.lower()):
                host = email.split('@')[1] if '@' in email else "zenika.com"
                prefix = email.split('@')[0] if '@' in email else "conflict"
                conflict_email = f"{prefix}.conflict.{random.randint(1000, 9999)}@{host}"
                new_u["email"] = conflict_email
                create_res = await http_client.post(f"{USERS_API_URL.rstrip('/')}/", json=new_u, headers=headers)
            if create_res.status_code >= 400:
                dur = int((time.monotonic() - t0) * 1000)
                _step_error("user_resolve", "Résolution & création d'identité", dur, f"Création utilisateur échouée (HTTP {create_res.status_code})")
                raise HTTPException(status_code=500, detail=f"User creation failed: {create_res.text}")
            user_id = create_res.json()["id"]

        dur = int((time.monotonic() - t0) * 1000)
        _step_ok("user_resolve", "Résolution & création d'identité", dur, f"User ID {user_id}")

        async def _bg_process_competencies_and_missions(bg_user_id, bg_structured_cv, bg_headers, bg_url):
            bg_errors = []
            new_competency_ids = []
            async with httpx.AsyncClient(timeout=120.0) as bg_http_client:
                try:
                    existing_comps = set()
                    try:
                        res = await bg_http_client.get(f"{COMPETENCIES_API_URL.rstrip('/')}/user/{bg_user_id}", headers=bg_headers, timeout=5.0)
                        if res.status_code == 200:
                            existing_comps = {c["id"] for c in res.json()}
                    except Exception as e:
                        logger.warning(f"Impossible de récupérer les compétences existantes: {e}")

                    sem = asyncio.Semaphore(3)

                    def normalize_comp(text):
                        if not text: return ""
                        text = text.strip().lower()
                        return "".join(c for c in unicodedata.normalize('NFKD', text) if unicodedata.category(c) != 'Mn')

                    async def resolve_comp_id(name: str) -> Optional[int]:
                        try:
                            res = await bg_http_client.get(
                                f"{COMPETENCIES_API_URL.rstrip('/')}/search",
                                params={"query": name, "limit": 5},
                                headers=bg_headers, timeout=5.0
                            )
                            if res.status_code == 200:
                                n_norm = normalize_comp(name)
                                for item in res.json().get("items", []):
                                    if normalize_comp(item.get("name", "")) == n_norm:
                                        return item["id"]
                                    aliases_raw = item.get("aliases") or ""
                                    for alias in aliases_raw.split(","):
                                        if normalize_comp(alias.strip()) == n_norm:
                                            return item["id"]
                        except Exception:
                            pass
                        return None

                    async def process_competency(comp):
                        async with sem:
                            name = sanitize_field(comp.get("name"))
                        if not name or not comp.get("practiced", True): return True
                        parent = sanitize_field(comp.get("parent"))
                        try:
                            c_id = await resolve_comp_id(name)
                            if not c_id:
                                p_id = None
                                if parent:
                                    p_id = await resolve_comp_id(parent)
                                    if not p_id:
                                        p_res = await bg_http_client.post(f"{COMPETENCIES_API_URL.rstrip('/')}/", json={"name": parent, "description": "Auto-identified from CV"}, headers=bg_headers)
                                        if p_res.status_code < 400: p_id = p_res.json()["id"]

                                aliases_str = ", ".join(comp.get("aliases", [])) if comp.get("aliases") else None
                                leaf_data = {"name": name, "description": "Candidate CV Skill", "aliases": aliases_str}
                                if p_id: leaf_data["parent_id"] = p_id

                                c_id = await resolve_comp_id(name)
                                if not c_id:
                                    c_res = await bg_http_client.post(f"{COMPETENCIES_API_URL.rstrip('/')}/", json=leaf_data, headers=bg_headers)
                                    if c_res.status_code < 400: c_id = c_res.json()["id"]

                            if c_id:
                                if c_id not in existing_comps:
                                    assign_res = await bg_http_client.post(f"{COMPETENCIES_API_URL.rstrip('/')}/user/{bg_user_id}/assign/{c_id}", headers=bg_headers)
                                    if assign_res.status_code >= 400: 
                                        return f"Échec d'assignation de '{name}'"
                                return True
                        except Exception as e:
                            return f"Erreur inattendue sur '{name}': {e}"
                        return f"Impossible de résoudre '{name}'"

                    comp_tasks = [process_competency(c) for c in bg_structured_cv.get("competencies", [])]
                    comp_results = await asyncio.gather(*comp_tasks, return_exceptions=True)
                    for res in comp_results:
                        if res is not True: bg_errors.append(str(res))
                except Exception as e:
                    bg_errors.append(f"Crash compétences: {e}")

                try:
                    missions_list = bg_structured_cv.get("missions", [])
                    cat_res = await bg_http_client.get(f"{ITEMS_API_URL.rstrip('/')}/categories", headers=bg_headers)
                    categories = cat_res.json() if cat_res.status_code == 200 else []

                    def find_cat_id(name):
                        for c in categories:
                            if c["name"].lower() == name.lower(): return c["id"]
                        return None

                    mission_cat_id = find_cat_id("Missions")
                    if not mission_cat_id:
                        m_res = await bg_http_client.post(f"{ITEMS_API_URL.rstrip('/')}/categories", json={"name": "Missions", "description": "Professional experiences"}, headers=bg_headers)
                        if m_res.status_code < 400: mission_cat_id = m_res.json()["id"]

                    sensitive_cat_id = find_cat_id("Restricted")
                    if not sensitive_cat_id:
                        s_res = await bg_http_client.post(f"{ITEMS_API_URL.rstrip('/')}/categories", json={"name": "Restricted", "description": "Sensitive missions"}, headers=bg_headers)
                        if s_res.status_code < 400: sensitive_cat_id = s_res.json()["id"]

                    item_data_list = []
                    for m in missions_list:
                        cat_ids = [mission_cat_id] if mission_cat_id else []
                        if m.get("is_sensitive") and sensitive_cat_id: cat_ids.append(sensitive_cat_id)
                        item_data_list.append({
                            "name": m["title"], "description": m.get("description", ""), "user_id": bg_user_id, "category_ids": cat_ids,
                            "metadata_json": {"company": m.get("company"), "competencies": m.get("competencies", []), "is_sensitive": m.get("is_sensitive", False), "start_date": m.get("start_date"), "end_date": m.get("end_date"), "duration": m.get("duration"), "mission_type": m.get("mission_type", "build"), "source": "CV Analysis"}
                        })

                    if item_data_list:
                        m_post = await bg_http_client.post(f"{ITEMS_API_URL.rstrip('/')}/bulk", json={"items": item_data_list}, headers=bg_headers)
                        if m_post.status_code >= 400: bg_errors.append("Création missions échouée")
                except Exception as e:
                    bg_errors.append(f"Crash missions: {e}")

            if bg_errors:
                try:
                    drive_api_url = os.getenv("DRIVE_API_URL", "http://drive_api:8006")
                    doc_id_match = re.search(r"/d/([a-zA-Z0-9_-]+)", bg_url)
                    if doc_id_match:
                        async with httpx.AsyncClient(timeout=10.0) as webhook_client:
                            await webhook_client.patch(f"{drive_api_url.rstrip('/')}/files/{doc_id_match.group(1)}", json={"status": "ERROR", "error_message": " | ".join(list(dict.fromkeys(bg_errors)))}, headers=bg_headers)
                except Exception:
                    pass

            try:
                async with httpx.AsyncClient(timeout=5.0) as _score_client:
                    _score_headers = dict(bg_headers)
                    from opentelemetry.propagate import inject as _inject
                    _inject(_score_headers)
                    await _score_client.post(f"{COMPETENCIES_API_URL.rstrip('/')}/evaluations/user/{bg_user_id}/ai-score-all?only_missing=true", headers=_score_headers, timeout=5.0)
            except Exception as e:
                logger.warning(f"Erreur globale lors du déclenchement du scoring IA: {e}")

        if background_tasks:
            background_tasks.add_task(_bg_process_competencies_and_missions, user_id, structured_cv, headers, url)
            _step_ok("competencies_missions", "Mapping RAG et Extraction", 0, "Délégué en Background Task (Asynchrone)")
            assigned_count = 0
        else:
            await _bg_process_competencies_and_missions(user_id, structured_cv, headers, url)
            assigned_count = 0

    # ── Étape 6 : Génération des embeddings vectoriels ────────────────────────
    t0 = time.monotonic()
    comp_keywords = [c.get("name") for c in structured_cv.get("competencies", []) if c.get("name")]
    distilled_content = _build_distilled_content(structured_cv)
    vector_data = None
    try:
        emb_res = await embed_content_with_retry(
            genai_client,
            model=os.getenv("GEMINI_EMBEDDING_MODEL"),
            contents=distilled_content
        )
        vector_data = emb_res.embeddings[0].values
        dur = int((time.monotonic() - t0) * 1000)
        _step_ok("embedding", "Génération des embeddings vectoriels", dur, f"{len(vector_data)} dimensions")
    except Exception as e:
        dur = int((time.monotonic() - t0) * 1000)
        _step_warn("embedding", "Génération des embeddings vectoriels", dur, f"Embedding échoué : {e}")
        pipeline_warnings.append("Embedding vectoriel échoué")

    # ── Étape 7 : Sauvegarde en base de données ───────────────────────────────
    t0 = time.monotonic()
    try:
        await db.execute(sa_delete(CVProfile).where(CVProfile.source_url == url))
        cv_record = CVProfile(
            user_id=user_id,
            source_url=url,
            source_tag=source_tag,
            extracted_competencies=structured_cv.get("competencies", []),
            current_role=structured_cv.get("current_role"),
            years_of_experience=structured_cv.get("years_of_experience"),
            summary=_coerce_to_str(structured_cv.get("summary")),
            competencies_keywords=comp_keywords,
            missions=structured_cv.get("missions", []),
            educations=structured_cv.get("educations", []),
            raw_content=raw_text,
            semantic_embedding=vector_data,
            imported_by_id=importer_id
        )
        db.add(cv_record)
        await db.commit()
        dur = int((time.monotonic() - t0) * 1000)
        _step_ok("db_save", "Sauvegarde en base de données", dur, f"CV ID utilisateur {user_id}")
    except Exception as e:
        dur = int((time.monotonic() - t0) * 1000)
        _step_error("db_save", "Sauvegarde en base de données", dur, str(e))
        CV_PROCESSING_TOTAL.labels(status="failure").inc()
        raise HTTPException(status_code=500, detail=f"Database save failed: {e}")

    CV_PROCESSING_TOTAL.labels(status="success").inc()

    response_data = {
        "status": "success",
        "message": "Import et analyse terminés avec succès",
        "user_id": user_id,
        "first_name": first_name,
        "last_name": last_name,
        "email": email,
        "is_anonymous": is_anonymous,
        "competencies_assigned": assigned_count,
        "warnings": pipeline_warnings,
        "steps": [s.model_dump() for s in pipeline_steps]
    }
    return CVResponse(**response_data)
