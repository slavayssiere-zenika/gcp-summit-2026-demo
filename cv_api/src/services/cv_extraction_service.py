import json
import logging
import os
import re
import urllib.parse
from datetime import datetime, timedelta, timezone
from typing import Optional

import httpx
from fastapi import HTTPException
from google.genai import types

from src.gemini_retry import generate_content_with_retry
from src.services.config import _CV_CACHE, COMPETENCIES_API_URL, PROMPTS_API_URL
from src.services.utils import _CV_RESPONSE_SCHEMA, build_taxonomy_context

logger = logging.getLogger(__name__)


class CVExtractionService:
    @staticmethod
    async def fetch_cv_content(
        url: str,
        google_token: Optional[str] = None,
        file_type: str = "google_doc",
    ) -> str:
        """Download the CV content."""
        parsed = urllib.parse.urlparse(url)
        hostname = parsed.hostname or ""

        if parsed.scheme not in ("http", "https"):
            raise HTTPException(status_code=400, detail="Invalid URL scheme")

        forbidden_hosts = ["localhost", "127.0.0.1", "0.0.0.0"]
        if hostname in forbidden_hosts or hostname.endswith(
                ".local") or hostname.endswith("_api"):
            raise HTTPException(
                status_code=400,
                detail="Internal URLs are not allowed")

        # ── Branche DOCX ─────────────────────────────────────────────────────
        if file_type == "docx":
            file_id_match = re.search(r"/file/d/([a-zA-Z0-9_-]+)", url)
            file_id = file_id_match.group(1) if file_id_match else None
            if not file_id:
                raise HTTPException(
                    status_code=400,
                    detail="DOCX URL invalide : impossible d'extraire le file_id depuis l'URL Drive.",
                )
            if not google_token:
                raise HTTPException(
                    status_code=400,
                    detail="DOCX Drive : google_access_token OAuth2 requis pour télécharger le fichier.",
                )
            download_url = f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"
            async with httpx.AsyncClient(timeout=60.0) as http_client:
                resp = await http_client.get(
                    download_url,
                    headers={"Authorization": f"Bearer {google_token}"},
                    follow_redirects=True,
                )
                if resp.status_code in (401, 403):
                    raise HTTPException(
                        status_code=400,
                        detail=(
                            f"Accès refusé pour le DOCX Drive (HTTP {
                                resp.status_code}). "
                            "Vérifiez les scopes OAuth2 du Service Account (drive.readonly minimum requis)."
                        ),
                    )
                resp.raise_for_status()

                # ── Validation MIME magic-bytes (sécurité P2 — anti-zip-bomb et fichiers maléficieux)
                _MAX_DOCX_SIZE = 10 * 1024 * 1024  # 10 Mo
                if len(resp.content) > _MAX_DOCX_SIZE:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Fichier DOCX trop volumineux ({
                            len(
                                resp.content) //
                            1024} Ko > 10 Mo max).",
                    )
                # Les fichiers DOCX sont des archives ZIP — magic bytes =
                # PK\x03\x04
                _DOCX_MAGIC = b"PK\x03\x04"
                if not resp.content[:4].startswith(_DOCX_MAGIC):
                    raise HTTPException(
                        status_code=400,
                        detail="Fichier refusé : le contenu ne correspond pas à un document DOCX valide (signature ZIP invalide).",
                    )

                from io import BytesIO
                import docx as python_docx
                doc = python_docx.Document(BytesIO(resp.content))

                # Extraction paragraphes (corps du CV)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Extraction tables (missions en tableau, expériences,
                # compétences tabulées)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(c.text.strip()
                                              for c in row.cells if c.text.strip())
                        if row_text:
                            paragraphs.append(row_text)

                text = "\n".join(paragraphs)
                if not text.strip():
                    raise HTTPException(
                        status_code=422,
                        detail="Le fichier DOCX est vide ou illisible (aucun texte extrait).",
                    )
                logger.info(
                    f"[_fetch_cv_content] DOCX extrait — {
                        len(text)} chars, file_id={file_id}")
                return text

        # ── Branche Google Doc natif ─────────────────────────────────────────
        if "docs.google.com/document/d/" in url:
            doc_id = re.search(r"/d/([a-zA-Z0-9_-]+)", url)
            if doc_id:
                file_id = doc_id.group(1)
                if not google_token:
                    logger.error(
                        f"[_fetch_cv_content] Aucun google_access_token pour le fichier Drive '{file_id}'. "
                        "L'API Drive v3 requiert un token OAuth2 valide (scope drive.readonly minimum)."
                    )
                    raise HTTPException(
                        status_code=400,
                        detail=(
                            "Document Google Docs privé : un google_access_token OAuth2 est requis. "
                            "Vérifiez que drive_api transmet bien le token dans le payload Pub/Sub."
                        ),
                    )
                export_url = f"https://www.googleapis.com/drive/v3/files/{file_id}/export"
                async with httpx.AsyncClient(timeout=30.0) as http_client:
                    resp = await http_client.get(
                        export_url,
                        params={"mimeType": "text/plain"},
                        headers={"Authorization": f"Bearer {google_token}"},
                        follow_redirects=True,
                    )
                    if resp.status_code in (401, 403):
                        raise HTTPException(
                            status_code=400,
                            detail=(
                                f"Accès refusé par l'API Drive (HTTP {
                                    resp.status_code}). "
                                "Vérifiez les scopes OAuth2 du Service Account "
                                "(drive.readonly ou drive.file minimum requis)."
                            ),
                        )
                    resp.raise_for_status()
                    return resp.text

        req_headers = {}
        if google_token:
            req_headers["Authorization"] = f"Bearer {google_token}"

        async with httpx.AsyncClient(timeout=30.0) as http_client:
            resp = await http_client.get(url, headers=req_headers, follow_redirects=True)
            if resp.status_code in (401, 403):
                raise HTTPException(
                    status_code=400,
                    detail="Accès refusé. Vérifiez que le document est accessible.",
                )
            resp.raise_for_status()
            return resp.text

    @staticmethod
    async def analyze_cv_with_llm(
        raw_text: str,
        headers: dict,
        genai_client,
    ) -> tuple[dict, object]:
        """
        Extrait les données structurées du CV via le modèle Gemini.
        Retourne un tuple: (structured_cv, response_usage_metadata)
        """
        prompt = None
        if _CV_CACHE["prompt"]["expires"] > datetime.now(
                timezone.utc) and _CV_CACHE["prompt"]["value"]:
            prompt = _CV_CACHE["prompt"]["value"]
        else:
            try:
                async with httpx.AsyncClient() as http_client:
                    res_prompt = await http_client.get(f"{PROMPTS_API_URL.rstrip('/')}/cv_api.extract_cv_info", headers=headers, timeout=5.0)
                    res_prompt.raise_for_status()
                    prompt = res_prompt.json()["value"]
                    _CV_CACHE["prompt"]["value"] = prompt
                    _CV_CACHE["prompt"]["expires"] = datetime.now(
                        timezone.utc) + timedelta(minutes=5)
            except Exception as e:
                if os.path.exists("cv_api.extract_cv_info.txt"):
                    with open("cv_api.extract_cv_info.txt", "r", encoding="utf-8") as f:
                        prompt = f.read()
                else:
                    raise HTTPException(
                        status_code=500, detail=f"Cannot fetch generic prompt: {e}")

        tree_context = ""
        if _CV_CACHE["tree_context"]["expires"] > datetime.now(
                timezone.utc) and _CV_CACHE["tree_context"]["value"]:
            tree_context = _CV_CACHE["tree_context"]["value"]
        else:
            try:
                async with httpx.AsyncClient() as http_client:
                    items: list = []
                    skip = 0
                    page_size = 100
                    while True:
                        page_res = await http_client.get(
                            f"{COMPETENCIES_API_URL.rstrip('/')}/",
                            params={"skip": skip, "limit": page_size},
                            headers=headers, timeout=5.0
                        )
                        if page_res.status_code != 200:
                            break
                        from shared.schemas.pagination import PaginationResponse
                        data = PaginationResponse[dict].model_validate(page_res.json())
                        page_items = data.items
                        items.extend(page_items)
                        if len(page_items) < page_size:
                            break
                        skip += page_size

                    if items:
                        tree_context, nb_parents, nb_leaves = build_taxonomy_context(
                            items)
                        _CV_CACHE["tree_context"]["value"] = tree_context
                        _CV_CACHE["tree_context"]["expires"] = datetime.now(
                            timezone.utc) + timedelta(minutes=5)
                        _CV_CACHE["tree_items"]["value"] = items
                        _CV_CACHE["tree_items"]["expires"] = datetime.now(
                            timezone.utc) + timedelta(minutes=5)
            except Exception as e:
                logger.warning(
                    f"Failed to fetch competencies tree for context: {e}")

        final_prompt = prompt + tree_context

        try:
            response = await generate_content_with_retry(
                genai_client,
                model=os.getenv("GEMINI_CV_MODEL", os.getenv("GEMINI_MODEL")),
                contents=[final_prompt, f"RESUME:\n{raw_text[:100000]}"],
                config=types.GenerateContentConfig(
                    response_mime_type="application/json",
                    response_schema=_CV_RESPONSE_SCHEMA,
                )
            )
            parsed_data = response.text
            structured_cv = json.loads(parsed_data)
            safe_meta = response.usage_metadata if hasattr(
                response, 'usage_metadata') else None
            return structured_cv, safe_meta
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"LLM Parsing failed: {e}")
